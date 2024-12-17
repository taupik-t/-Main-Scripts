# import pytesseract
# from PIL import Image
# from docx import Document
# from docx.shared import Pt

# def ocr_to_word(input_image_path, output_word_path):
#     # Perform OCR on the input image
#     extracted_text = pytesseract.image_to_string(Image.open(input_image_path))

#     # Create a new Word document
#     doc = Document()
    
#     # Add the extracted text to the Word document
#     paragraph = doc.add_paragraph()
#     run = paragraph.add_run(extracted_text)
#     font = run.font
#     font.size = Pt(12)  # Set the font size (you can change it as needed)

#     # Save the Word document
#     doc.save(output_word_path)

# if __name__ == "__main__":
#     input_image_path = "img.jpg"  # Replace with the path to your image
#     output_word_path = "document.docx"  # Replace with the desired output path

#     ocr_to_word(input_image_path, output_word_path)


import pytesseract
from PIL import Image
from docx import Document
from docx.shared import Pt

def ocr_to_word(input_image_path, output_word_path):
    # Set the language to Indonesian
    lang = 'ind'

    # Perform OCR on the input image with the specified language
    extracted_text = pytesseract.image_to_string(Image.open(input_image_path), lang=lang)

    # Create a new Word document
    doc = Document()
    
    # Add the extracted text to the Word document
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(extracted_text)
    font = run.font
    font.size = Pt(12)  # Set the font size (you can change it as needed)

    # Save the Word document
    doc.save(output_word_path)

if __name__ == "__main__":
    input_image_path = "img.jpg"  # Replace with the path to your image
    output_word_path = "document1.docx"  # Replace with the desired output path

    ocr_to_word(input_image_path, output_word_path)
